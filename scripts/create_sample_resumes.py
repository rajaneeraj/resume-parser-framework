"""
Script to generate synthetic sample resumes for testing.

Creates a PDF and a DOCX resume with realistic content in the
'samples/' directory. These files can be used to test the framework.

Usage:
    python scripts/create_sample_resumes.py
"""

import os
import sys
from pathlib import Path

# Add the src directory to the path so we can import project modules if needed
PROJECT_ROOT = Path(__file__).resolve().parent.parent
SAMPLES_DIR = PROJECT_ROOT / "samples"


def create_sample_docx(output_path: Path) -> None:
    """Create a synthetic Word resume.

    Args:
        output_path: Path where the .docx file will be saved.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Header: Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run("Jane Doe")
    name_run.bold = True
    name_run.font.size = Pt(24)

    # --- Contact Info ---
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run("jane.doe@gmail.com | (555) 123-4567 | San Francisco, CA")

    # --- Summary ---
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Experienced Machine Learning Engineer with 5+ years of expertise in "
        "designing and deploying production ML systems. Passionate about "
        "Natural Language Processing and Large Language Models."
    )

    # --- Experience ---
    doc.add_heading("Experience", level=2)

    doc.add_heading("Senior ML Engineer — TechCorp Inc.", level=3)
    doc.add_paragraph("January 2021 – Present")
    doc.add_paragraph(
        "Led development of an NLP pipeline processing 10M+ documents daily "
        "using Python, TensorFlow, and AWS infrastructure.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Built and deployed LLM-based summarization service using "
        "Docker and Kubernetes on GCP.",
        style="List Bullet",
    )

    doc.add_heading("ML Engineer — DataSolutions Ltd.", level=3)
    doc.add_paragraph("June 2018 – December 2020")
    doc.add_paragraph(
        "Developed classification models achieving 95% accuracy using "
        "Scikit-learn and PyTorch.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Implemented REST API endpoints with Flask for model serving.",
        style="List Bullet",
    )

    # --- Education ---
    doc.add_heading("Education", level=2)
    doc.add_paragraph("M.S. Computer Science — Stanford University, 2018")
    doc.add_paragraph("B.S. Mathematics — UC Berkeley, 2016")

    # --- Skills ---
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(
        "Python, Machine Learning, Deep Learning, TensorFlow, PyTorch, "
        "NLP, LLM, AWS, GCP, Docker, Kubernetes, SQL, Git, REST API, "
        "Scikit-learn, Pandas, NumPy, Flask, Data Analysis"
    )

    doc.save(str(output_path))
    print(f"  [OK] Created: {output_path}")


def create_sample_pdf(output_path: Path) -> None:
    """Create a synthetic PDF resume.

    Args:
        output_path: Path where the .pdf file will be saved.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    name_style = ParagraphStyle(
        "NameStyle",
        parent=styles["Title"],
        fontSize=22,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    contact_style = ParagraphStyle(
        "ContactStyle",
        parent=styles["Normal"],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = styles["Heading2"]
    body_style = styles["Normal"]
    bullet_style = ParagraphStyle(
        "BulletStyle",
        parent=styles["Normal"],
        leftIndent=20,
        spaceAfter=4,
    )

    elements: list = []

    # --- Header ---
    elements.append(Paragraph("John Smith", name_style))
    elements.append(
        Paragraph(
            "john.smith@outlook.com | (555) 987-6543 | New York, NY",
            contact_style,
        )
    )
    elements.append(Spacer(1, 12))

    # --- Summary ---
    elements.append(Paragraph("Professional Summary", heading_style))
    elements.append(
        Paragraph(
            "Full-stack Software Engineer with 7 years of experience building "
            "scalable web applications. Strong background in React, Node.js, "
            "and cloud-native architectures. Experienced with Agile methodologies.",
            body_style,
        )
    )
    elements.append(Spacer(1, 8))

    # --- Experience ---
    elements.append(Paragraph("Experience", heading_style))

    elements.append(Paragraph("<b>Lead Software Engineer — WebScale Inc.</b>", body_style))
    elements.append(Paragraph("March 2020 – Present", body_style))
    elements.append(
        Paragraph(
            "• Architected microservices backend using Node.js, Express.js, "
            "and PostgreSQL handling 50K+ requests/sec.",
            bullet_style,
        )
    )
    elements.append(
        Paragraph(
            "• Led migration of monolith to Kubernetes on AWS, reducing "
            "infrastructure costs by 40%.",
            bullet_style,
        )
    )
    elements.append(Spacer(1, 6))

    elements.append(
        Paragraph("<b>Software Engineer — AppDev Corp.</b>", body_style)
    )
    elements.append(Paragraph("July 2016 – February 2020", body_style))
    elements.append(
        Paragraph(
            "• Built responsive frontend applications using React, TypeScript, "
            "and Next.js.",
            bullet_style,
        )
    )
    elements.append(
        Paragraph(
            "• Implemented CI/CD pipelines with Jenkins and Docker for "
            "automated deployments.",
            bullet_style,
        )
    )
    elements.append(Spacer(1, 8))

    # --- Education ---
    elements.append(Paragraph("Education", heading_style))
    elements.append(
        Paragraph("B.S. Computer Science — MIT, 2016", body_style)
    )
    elements.append(Spacer(1, 8))

    # --- Skills ---
    elements.append(Paragraph("Skills", heading_style))
    elements.append(
        Paragraph(
            "JavaScript, TypeScript, Python, React, Node.js, Express.js, "
            "Next.js, PostgreSQL, MongoDB, AWS, Docker, Kubernetes, "
            "CI/CD, Git, REST API, GraphQL, Microservices, Agile, Scrum",
            body_style,
        )
    )

    doc.build(elements)
    print(f"  [OK] Created: {output_path}")


def main() -> None:
    """Generate sample resume files."""
    print("Generating sample resumes...")

    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

    create_sample_docx(SAMPLES_DIR / "jane_doe_resume.docx")
    create_sample_pdf(SAMPLES_DIR / "john_smith_resume.pdf")

    print(f"\nSample resumes saved to: {SAMPLES_DIR}")


if __name__ == "__main__":
    main()
