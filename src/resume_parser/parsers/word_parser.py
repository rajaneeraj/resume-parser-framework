"""
WordParser — Extracts raw text content from Word (.docx) files.

Uses python-docx to read paragraphs and table cells from a
Word document, producing a single text string.
"""

import logging
from pathlib import Path
from typing import ClassVar

from docx import Document

from resume_parser.parsers.base import FileParser

logger = logging.getLogger(__name__)


class WordParser(FileParser):
    """Concrete parser for Word Document (.docx) resume files.

    Reads all paragraphs and table cell contents from a .docx file
    and returns the combined text.
    """

    supported_extensions: ClassVar[set[str]] = {".docx"}

    def _extract_text(self, path: Path) -> str:
        """Extract text from paragraphs and tables of a Word document.

        Args:
            path: Path to the .docx file.

        Returns:
            Combined text from paragraphs and tables, separated by newlines.

        Raises:
            RuntimeError: If the document cannot be read or is corrupted.
        """
        try:
            doc = Document(str(path))
            parts: list[str] = []

            # Word XML namespace for paragraph and text-run elements
            _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Extract text from text boxes (w:txbxContent).
            # Many formatted resumes place the candidate name and contact
            # info inside Word text boxes (drawing objects).  python-docx's
            # doc.paragraphs only covers top-level body paragraphs, so we
            # need to reach into the XML to find text box content.
            # A set is used to de-duplicate because MC:AlternateContent
            # blocks often duplicate the same text box content.
            seen_textbox_lines: set[str] = set()
            for txbx_para in doc.element.body.findall(
                f".//{{{_W_NS}}}txbxContent/{{{_W_NS}}}p"
            ):
                runs = txbx_para.findall(f".//{{{_W_NS}}}t")
                line = "".join(r.text or "" for r in runs).strip()
                if line and line not in seen_textbox_lines:
                    seen_textbox_lines.add(line)
                    parts.append(line)

            if seen_textbox_lines:
                logger.debug(
                    "Extracted %d text-box lines from %s",
                    len(seen_textbox_lines),
                    path.name,
                )

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)

            # Extract text from tables (resumes sometimes use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            logger.debug(
                "Extracted %d text segments from %s", len(parts), path.name
            )
            return "\n".join(parts)

        except Exception as exc:
            raise RuntimeError(
                f"Failed to extract text from Word document '{path.name}': {exc}"
            ) from exc
