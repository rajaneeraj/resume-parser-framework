"""
Shared pytest fixtures and helpers for the test suite.

Provides:
- Temporary file fixtures for PDF and DOCX generation.
- Path helpers for the samples directory.

Sample text constants are in tests/sample_data.py for clean imports.
"""

import sys
from pathlib import Path

import pytest

# Ensure the src directory is on the Python path for test imports
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "src"))


# ──────────────────────────────────────────────────────────────
# Fixtures for creating temporary test files
# ──────────────────────────────────────────────────────────────


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """Create a temporary .docx file with sample resume content."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@gmail.com | (555) 123-4567")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, Machine Learning, AWS, Docker, TensorFlow")

    file_path = tmp_path / "test_resume.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def tmp_pdf(tmp_path: Path) -> Path:
    """Create a temporary .pdf file with sample resume content."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet

    file_path = tmp_path / "test_resume.pdf"
    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    styles = getSampleStyleSheet()

    elements = [
        Paragraph("John Smith", styles["Title"]),
        Paragraph("john.smith@outlook.com | (555) 987-6543", styles["Normal"]),
        Paragraph("Skills", styles["Heading2"]),
        Paragraph(
            "JavaScript, TypeScript, React, Node.js, AWS, Docker",
            styles["Normal"],
        ),
    ]
    doc.build(elements)
    return file_path


@pytest.fixture
def tmp_empty_docx(tmp_path: Path) -> Path:
    """Create a temporary empty .docx file."""
    from docx import Document

    doc = Document()
    file_path = tmp_path / "empty_resume.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def tmp_txt(tmp_path: Path) -> Path:
    """Create a temporary .txt file (unsupported format)."""
    file_path = tmp_path / "resume.txt"
    file_path.write_text("This is plain text, not a supported format.")
    return file_path
