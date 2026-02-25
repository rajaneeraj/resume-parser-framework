"""
Unit tests for the parse_resumes.py CLI module.

Tests argument parsing, folder scanning, output writing,
and archive logic using temporary directories and mocks.
"""

import json
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

# We need to import from the CLI module — it's at the project root,
# so we add the project root to sys.path via conftest.py (already done).
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from parse_resumes import (
    parse_args,
    discover_resumes,
    build_name_extractor,
    build_skills_extractor,
    archive_file,
    run,
)


# ──────────────────────────────────────────────────────────────
# Argument parsing tests
# ──────────────────────────────────────────────────────────────


class TestParseArgs:
    """Tests for CLI argument parsing."""

    def test_default_values(self):
        """Should use default values when no args are provided."""
        args = parse_args([])
        assert args.input_dir == Path("resumes")
        assert args.output_dir == Path("output")
        assert args.archive_dir == Path("archive")
        assert args.no_archive is False
        assert args.no_llm is False

    def test_custom_input_dir(self):
        """Should accept custom input directory."""
        args = parse_args(["--input-dir", "my_resumes"])
        assert args.input_dir == Path("my_resumes")

    def test_custom_output_dir(self):
        """Should accept custom output directory."""
        args = parse_args(["--output-dir", "my_output"])
        assert args.output_dir == Path("my_output")

    def test_custom_archive_dir(self):
        """Should accept custom archive directory."""
        args = parse_args(["--archive-dir", "my_archive"])
        assert args.archive_dir == Path("my_archive")

    def test_no_archive_flag(self):
        """Should set no_archive to True when flag is passed."""
        args = parse_args(["--no-archive"])
        assert args.no_archive is True

    def test_no_llm_flag(self):
        """Should set no_llm to True when flag is passed."""
        args = parse_args(["--no-llm"])
        assert args.no_llm is True

    def test_all_flags_combined(self):
        """Should handle all flags together."""
        args = parse_args([
            "--input-dir", "in",
            "--output-dir", "out",
            "--archive-dir", "arch",
            "--no-archive",
            "--no-llm",
        ])
        assert args.input_dir == Path("in")
        assert args.output_dir == Path("out")
        assert args.archive_dir == Path("arch")
        assert args.no_archive is True
        assert args.no_llm is True


# ──────────────────────────────────────────────────────────────
# Folder scanning tests
# ──────────────────────────────────────────────────────────────


class TestDiscoverResumes:
    """Tests for the discover_resumes function."""

    def test_finds_pdf_and_docx(self, tmp_path: Path):
        """Should find .pdf and .docx files."""
        (tmp_path / "resume1.pdf").write_text("pdf content")
        (tmp_path / "resume2.docx").write_text("docx content")

        result = discover_resumes(tmp_path)
        names = [p.name for p in result]

        assert len(result) == 2
        assert "resume1.pdf" in names
        assert "resume2.docx" in names

    def test_ignores_other_extensions(self, tmp_path: Path):
        """Should ignore .txt, .jpg, and other non-resume files."""
        (tmp_path / "resume.pdf").write_text("pdf content")
        (tmp_path / "notes.txt").write_text("text notes")
        (tmp_path / "photo.jpg").write_bytes(b"image")

        result = discover_resumes(tmp_path)
        assert len(result) == 1
        assert result[0].name == "resume.pdf"

    def test_recursive_scan(self, tmp_path: Path):
        """Should recursively find files in subdirectories."""
        sub = tmp_path / "subdir"
        sub.mkdir()
        (sub / "nested_resume.docx").write_text("docx content")
        (tmp_path / "top_resume.pdf").write_text("pdf content")

        result = discover_resumes(tmp_path)
        assert len(result) == 2

    def test_empty_directory(self, tmp_path: Path):
        """Should return empty list for directory with no resumes."""
        result = discover_resumes(tmp_path)
        assert result == []

    def test_nonexistent_directory(self, tmp_path: Path):
        """Should return empty list for non-existent directory."""
        result = discover_resumes(tmp_path / "doesnt_exist")
        assert result == []

    def test_sorted_output(self, tmp_path: Path):
        """Should return results in sorted order."""
        (tmp_path / "c_resume.pdf").write_text("c")
        (tmp_path / "a_resume.pdf").write_text("a")
        (tmp_path / "b_resume.docx").write_text("b")

        result = discover_resumes(tmp_path)
        names = [p.name for p in result]
        assert names == sorted(names)


# ──────────────────────────────────────────────────────────────
# Strategy fallback tests
# ──────────────────────────────────────────────────────────────


class TestBuildNameExtractor:
    """Tests for name extractor strategy selection."""

    def test_falls_back_to_rule_based(self):
        """Should fall back to RuleBasedNameExtractor when spaCy is unavailable."""
        with patch(
            "resume_parser.extractors.spacy_name_extractor.SpacyNameExtractor",
            side_effect=ImportError("no spacy"),
        ):
            from resume_parser.extractors import RuleBasedNameExtractor

            extractor = build_name_extractor()
            assert isinstance(extractor, RuleBasedNameExtractor)


class TestBuildSkillsExtractor:
    """Tests for skills extractor strategy selection."""

    def test_no_llm_flag_forces_keyword(self):
        """Should use KeywordSkillsExtractor when --no-llm is set."""
        from resume_parser.extractors import KeywordSkillsExtractor

        extractor = build_skills_extractor(no_llm=True)
        assert isinstance(extractor, KeywordSkillsExtractor)

    def test_no_api_key_uses_keyword(self):
        """Should use keyword extractor when GEMINI_API_KEY is not set."""
        from resume_parser.extractors import KeywordSkillsExtractor

        with patch.dict("os.environ", {}, clear=True):
            extractor = build_skills_extractor(no_llm=False)
            assert isinstance(extractor, KeywordSkillsExtractor)

    def test_placeholder_key_uses_keyword(self):
        """Should treat 'your_api_key_here' as no key."""
        from resume_parser.extractors import KeywordSkillsExtractor

        with patch.dict("os.environ", {"GEMINI_API_KEY": "your_api_key_here"}):
            extractor = build_skills_extractor(no_llm=False)
            assert isinstance(extractor, KeywordSkillsExtractor)


# ──────────────────────────────────────────────────────────────
# Archive tests
# ──────────────────────────────────────────────────────────────


class TestArchiveFile:
    """Tests for the archive_file function."""

    def test_moves_file_to_timestamped_dir(self, tmp_path: Path):
        """Should move the file to archive/<timestamp>/filename."""
        input_dir = tmp_path / "input"
        input_dir.mkdir()
        file = input_dir / "resume.pdf"
        file.write_text("pdf content")

        archive_dir = tmp_path / "archive"

        result = archive_file(file, archive_dir, "2026-02-25_120000", input_dir)

        assert result.exists()
        assert not file.exists()
        assert "2026-02-25_120000" in str(result)
        assert result.name == "resume.pdf"

    def test_preserves_subdirectory_structure(self, tmp_path: Path):
        """Should preserve relative path structure in archive."""
        input_dir = tmp_path / "input"
        sub = input_dir / "subdir"
        sub.mkdir(parents=True)
        file = sub / "nested.docx"
        file.write_text("content")

        archive_dir = tmp_path / "archive"

        result = archive_file(file, archive_dir, "2026-02-25_120000", input_dir)

        assert "subdir" in str(result)
        assert result.name == "nested.docx"


# ──────────────────────────────────────────────────────────────
# Full run integration tests
# ──────────────────────────────────────────────────────────────


class TestRunCLI:
    """Integration tests for the CLI run function."""

    def _create_test_docx(self, path: Path) -> None:
        """Create a minimal test .docx file."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test User")
        doc.add_paragraph("test@example.com")
        doc.add_paragraph("Skills: Python, Java")
        doc.save(str(path))

    def test_run_processes_files(self, tmp_path: Path):
        """Should process resume files and write results.json."""
        input_dir = tmp_path / "resumes"
        input_dir.mkdir()
        self._create_test_docx(input_dir / "test.docx")

        output_dir = tmp_path / "output"
        archive_dir = tmp_path / "archive"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--archive-dir", str(archive_dir),
            "--no-llm",
        ])

        exit_code = run(args)

        assert exit_code == 0
        assert (output_dir / "results.json").exists()

        results = json.loads((output_dir / "results.json").read_text())
        assert len(results) == 1
        assert results[0]["source_file"] == "test.docx"
        assert results[0]["email"] == "test@example.com"

    def test_run_no_files_returns_2(self, tmp_path: Path):
        """Should return exit code 2 when no resume files are found."""
        input_dir = tmp_path / "empty_resumes"
        input_dir.mkdir()
        output_dir = tmp_path / "output"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--no-llm",
        ])

        exit_code = run(args)
        assert exit_code == 2

    def test_run_archives_processed_files(self, tmp_path: Path):
        """Should move processed files to archive directory."""
        input_dir = tmp_path / "resumes"
        input_dir.mkdir()
        self._create_test_docx(input_dir / "test.docx")

        output_dir = tmp_path / "output"
        archive_dir = tmp_path / "archive"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--archive-dir", str(archive_dir),
            "--no-llm",
        ])

        run(args)

        # Original file should be moved
        assert not (input_dir / "test.docx").exists()
        # Archive should have the file
        archived = list(archive_dir.rglob("test.docx"))
        assert len(archived) == 1

    def test_run_no_archive_flag(self, tmp_path: Path):
        """Should not move files when --no-archive is set."""
        input_dir = tmp_path / "resumes"
        input_dir.mkdir()
        self._create_test_docx(input_dir / "test.docx")

        output_dir = tmp_path / "output"
        archive_dir = tmp_path / "archive"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--archive-dir", str(archive_dir),
            "--no-archive",
            "--no-llm",
        ])

        run(args)

        # File should still be in input dir
        assert (input_dir / "test.docx").exists()
        # Archive should not exist
        assert not archive_dir.exists()

    def test_run_writes_errors_on_failure(self, tmp_path: Path):
        """Should write errors.json when a file fails to parse."""
        input_dir = tmp_path / "resumes"
        input_dir.mkdir()
        # Create a corrupt file that will fail to parse
        corrupt = input_dir / "bad.pdf"
        corrupt.write_text("this is not a real pdf")

        output_dir = tmp_path / "output"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--no-archive",
            "--no-llm",
        ])

        exit_code = run(args)

        assert exit_code == 1  # Partial failure
        assert (output_dir / "errors.json").exists()
        errors = json.loads((output_dir / "errors.json").read_text())
        assert len(errors) == 1
        assert "bad.pdf" in errors[0]["file"]

    def test_batch_continues_after_failure(self, tmp_path: Path):
        """Should continue processing remaining files after a failure."""
        input_dir = tmp_path / "resumes"
        input_dir.mkdir()

        # Create one corrupt file and one valid file
        (input_dir / "bad.pdf").write_text("not a pdf")
        self._create_test_docx(input_dir / "good.docx")

        output_dir = tmp_path / "output"

        args = parse_args([
            "--input-dir", str(input_dir),
            "--output-dir", str(output_dir),
            "--no-archive",
            "--no-llm",
        ])

        exit_code = run(args)

        assert exit_code == 1  # Partial failure
        results = json.loads((output_dir / "results.json").read_text())
        errors = json.loads((output_dir / "errors.json").read_text())

        # One success, one failure
        assert len(results) == 1
        assert len(errors) == 1
        assert results[0]["source_file"] == "good.docx"
