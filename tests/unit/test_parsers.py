"""
Unit tests for FileParser, PDFParser, and WordParser.
"""

from pathlib import Path

import pytest

from resume_parser.parsers.base import FileParser
from resume_parser.parsers.pdf_parser import PDFParser
from resume_parser.parsers.word_parser import WordParser

# ──────────────────────────────────────────────────────────────
# FileParser ABC tests
# ──────────────────────────────────────────────────────────────


class TestFileParserAbstract:
    """Test that FileParser cannot be instantiated directly."""

    def test_cannot_instantiate_abstract_class(self):
        """FileParser is abstract and should not be directly instantiable."""
        with pytest.raises(TypeError):
            FileParser()


# ──────────────────────────────────────────────────────────────
# PDFParser tests
# ──────────────────────────────────────────────────────────────


class TestPDFParser:
    """Tests for PDFParser."""

    def test_parse_valid_pdf(self, tmp_pdf: Path):
        """Should extract text from a valid PDF file."""
        parser = PDFParser()
        text = parser.parse(str(tmp_pdf))

        assert isinstance(text, str)
        assert len(text) > 0
        assert "John Smith" in text

    def test_parse_extracts_email(self, tmp_pdf: Path):
        """Should extract email address from PDF content."""
        parser = PDFParser()
        text = parser.parse(str(tmp_pdf))

        assert "john.smith@outlook.com" in text

    def test_parse_extracts_skills(self, tmp_pdf: Path):
        """Should extract skills section from PDF content."""
        parser = PDFParser()
        text = parser.parse(str(tmp_pdf))

        assert "JavaScript" in text

    def test_parse_file_not_found(self):
        """Should raise FileNotFoundError for non-existent file."""
        parser = PDFParser()
        with pytest.raises(FileNotFoundError, match="File not found"):
            parser.parse("/nonexistent/path/resume.pdf")

    def test_parse_wrong_extension(self, tmp_docx: Path):
        """Should raise ValueError for wrong file extension."""
        parser = PDFParser()
        with pytest.raises(ValueError, match="Unsupported file extension"):
            parser.parse(str(tmp_docx))

    def test_parse_directory_not_file(self, tmp_path: Path):
        """Should raise ValueError when given a directory path."""
        parser = PDFParser()
        with pytest.raises(ValueError, match="not a file"):
            parser.parse(str(tmp_path))

    def test_supported_extensions(self):
        """PDFParser should only support .pdf extension."""
        parser = PDFParser()
        assert parser.supported_extensions == {".pdf"}


# ──────────────────────────────────────────────────────────────
# WordParser tests
# ──────────────────────────────────────────────────────────────


class TestWordParser:
    """Tests for WordParser."""

    def test_parse_valid_docx(self, tmp_docx: Path):
        """Should extract text from a valid .docx file."""
        parser = WordParser()
        text = parser.parse(str(tmp_docx))

        assert isinstance(text, str)
        assert len(text) > 0
        assert "Jane Doe" in text

    def test_parse_extracts_email(self, tmp_docx: Path):
        """Should extract email address from DOCX content."""
        parser = WordParser()
        text = parser.parse(str(tmp_docx))

        assert "jane.doe@gmail.com" in text

    def test_parse_extracts_skills(self, tmp_docx: Path):
        """Should extract skills section from DOCX content."""
        parser = WordParser()
        text = parser.parse(str(tmp_docx))

        assert "Python" in text

    def test_parse_empty_docx(self, tmp_empty_docx: Path):
        """Should handle an empty DOCX without errors."""
        parser = WordParser()
        text = parser.parse(str(tmp_empty_docx))

        assert isinstance(text, str)
        # Empty document should produce empty or whitespace-only text
        assert text.strip() == ""

    def test_parse_file_not_found(self):
        """Should raise FileNotFoundError for non-existent file."""
        parser = WordParser()
        with pytest.raises(FileNotFoundError, match="File not found"):
            parser.parse("/nonexistent/path/resume.docx")

    def test_parse_wrong_extension(self, tmp_pdf: Path):
        """Should raise ValueError for wrong file extension."""
        parser = WordParser()
        with pytest.raises(ValueError, match="Unsupported file extension"):
            parser.parse(str(tmp_pdf))

    def test_supported_extensions(self):
        """WordParser should only support .docx extension."""
        parser = WordParser()
        assert parser.supported_extensions == {".docx"}

    def test_parse_docx_with_table(self, tmp_path: Path):
        """Should extract text from tables within a DOCX file."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"

        file_path = tmp_path / "table_resume.docx"
        doc.save(str(file_path))

        parser = WordParser()
        text = parser.parse(str(file_path))

        assert "Python" in text
        assert "Expert" in text

    def test_parse_docx_with_textbox(self, tmp_path: Path):
        """Should extract text from Word text boxes (w:txbxContent)."""
        from docx import Document
        from lxml import etree

        doc = Document()
        # Manually inject a text box into the document XML
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        wpc_ns = "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
        txbx_xml = f"""
        <w:p xmlns:w="{w_ns}">
          <w:r>
            <w:rPr/>
            <w:drawing>
              <wpc:wpc xmlns:wpc="{wpc_ns}">
                <w:txbxContent xmlns:w="{w_ns}">
                  <w:p>
                    <w:r><w:t>Neeraj Raja</w:t></w:r>
                  </w:p>
                  <w:p>
                    <w:r><w:t>neeraj.raja@hotmail.com</w:t></w:r>
                  </w:p>
                </w:txbxContent>
              </wpc:wpc>
            </w:drawing>
          </w:r>
        </w:p>
        """
        txbx_element = etree.fromstring(txbx_xml)
        # Insert the text box element at the beginning of the body
        doc.element.body.insert(0, txbx_element)

        # Also add a regular paragraph
        doc.add_paragraph("PROFESSIONAL SUMMARY")

        file_path = tmp_path / "textbox_resume.docx"
        doc.save(str(file_path))

        parser = WordParser()
        text = parser.parse(str(file_path))

        # Text box content should appear in extracted text
        assert "Neeraj Raja" in text
        assert "neeraj.raja@hotmail.com" in text
        # Regular paragraph text should also be there
        assert "PROFESSIONAL SUMMARY" in text
        # Text box content should appear BEFORE paragraph content
        neeraj_pos = text.index("Neeraj Raja")
        summary_pos = text.index("PROFESSIONAL SUMMARY")
        assert neeraj_pos < summary_pos
